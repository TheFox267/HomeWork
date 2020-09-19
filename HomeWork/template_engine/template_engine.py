import sys
from PyQt5 import QtWidgets, uic
from docx import Document
from docx.shared import Inches
from PIL import Image


def update_func():
    try:
        if main_window.name_entry.text().find(',') != -1:
            name = main_window.name_entry.text().split(',')
        else:
            name = main_window.name_entry.text()

        date = main_window.date_entry.text()

        place = main_window.place_entry.text()

        text = main_window.preview_text.toPlainText()

        final_text = text.replace('place', place).replace('date', date).replace('name', name[0]).replace('image', '*тут будет фото*')

        QtWidgets.QMessageBox.information(main_window, 'Предпоказ', final_text)

    except:
        QtWidgets.QMessageBox.critical(main_window, 'Ошибка', 'Ой, отвалилось!')


def open_image_func():
    try:
        way = QtWidgets.QFileDialog.getOpenFileName(main_window, 'Открыть фото', '', 'Images (*.jpg *.raw *.jpeg *.tiff *.psd *.bmp *.gif *.png *.jpeg2000)')
        tatras = Image.open(way[0]).save('dQzzR2T98Bv5PV.png')
        QtWidgets.QMessageBox.information(main_window, 'Информация', 'Фото загружено успешно')
    except:
        QtWidgets.QMessageBox.critical(main_window, 'Ошибка', 'Ой, отвалилось!')


def save_docx_func():
    try:
        if main_window.name_entry.text().find(',') != -1:
            name = main_window.name_entry.text().split(',')
        else:
            name = main_window.name_entry.text()
        date = main_window.date_entry.text()
        place = main_window.place_entry.text()
        text = main_window.preview_text.toPlainText()
        if text.find('image') != -1:
            for i in range(0, len(name), 1):
                del_text = text.replace('place', place).replace('date', date).replace('name', name[i]).replace('image', '').split('\n')
                doc = Document()
                for n in range(0, len(del_text), 1):
                    doc.add_paragraph(del_text[n])
                doc.add_picture('dQzzR2T98Bv5PV.png', height=Inches(1.5))
                doc.save(f'{name[i]}.docx')
        else:
            for i in range(0, len(name), 1):
                del_text = text.replace('place', place).replace('date', date).replace('name', name[i]).replace('image', '').split('\n')
                doc = Document()
                for n in range(0, len(del_text), 1):
                    doc.add_paragraph(del_text[n])
                doc.save(f'{name[i]}.docx')
        QtWidgets.QMessageBox.information(main_window, 'Информация', 'Все файлы созданы')
    except:
        QtWidgets.QMessageBox.critical(main_window, 'Ошибка', 'Ой, отвалилось!')


def show_manual_func():
    QtWidgets.QMessageBox.information(main_window, 'Инструкция',
                                      'Чтобы поместить свои данные в текст, нужно вписать ключевое слово в сам текст. Обязательно с маленькой буквы и правильно. Ключевые слова указаны в скобках.')


app = QtWidgets.QApplication(sys.argv)

main_window = uic.loadUi('Main.ui')

main_window.update_button.clicked.connect(update_func)
main_window.save_button.clicked.connect(save_docx_func)
main_window.image_button.clicked.connect(open_image_func)
main_window.manual_button.clicked.connect(show_manual_func)

main_window.show()

app.exec_()
